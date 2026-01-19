import os
import shutil
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client

class PDFMerger:
    # Mapeo de claves de checkbox a nombres de archivo físicos
    SECTION_MAP = {
        "portadas": "Portadas.pdf",
        # "contenido_separadores": GENERATED DINAMICAMENTE
        # "carta_presentacion": "carta de presentacion.pdf", # Se elimina o se maneja aparte
        "paginas_estandar": "paginas_estandar.pdf",
        "cuadro_experiencia": "experiencia_fotografico.pdf",
        "certificados_trabajos": "certificaciones.pdf",
        "seguridad_alturas": "seguridad_alturas.pdf",
        "programa_prevencion": "resumen_alturas.pdf",
        "sgsst_certificado": "ministerio.pdf",
        "documentacion_legal": "antecedentes.pdf",
        # "anexos": "agradecimiento.pdf" # Verificaremos si es este
    }

    def __init__(self, templates_dir):
        self.templates_dir = templates_dir

    def generate_separators_pdf(self, selected_sections_keys, output_path, page_mapping=None):
        """
        Genera un PDF con la tabla de contenido (Separadores) basado en las secciones seleccionadas.
        Crea un Word temporal y lo convierte a PDF.
        :param page_mapping: Dict mapping section_key -> start_page_number (int)
        """
        try:
            doc = Document()
            
            # Estilos básicos
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)

            # Título
            title = doc.add_paragraph('TABLA DE CONTENIDO')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.bold = True
            title_run.font.size = Pt(16)
            doc.add_paragraph() # Espacio

            # Lista de secciones (A, B, C...)
            # Mapa de nombres legibles para la tabla de contenido
            readable_names = {
                "portadas": "PORTADAS",
                "carta_presentacion": "CARTA DE PRESENTACIÓN",
                "paginas_estandar": "PÓLIZAS Y PERSONAL",
                "cuadro_experiencia": "EXPERIENCIA ESPECÍFICA",
                "certificados_trabajos": "CERTIFICACIONES",
                "seguridad_alturas": "SEGURIDAD EN ALTURAS",
                "programa_prevencion": "PROGRAMA DE PREVENCIÓN",
                "sgsst_certificado": "SISTEMA DE GESTIÓN (SGSST)",
                "presupuesto_programacion": "PRESUPUESTO Y PROGRAMACIÓN",
                "documentacion_legal": "DOCUMENTACIÓN LEGAL",
                "anexos": "ANEXOS",
                "propuesta_tecnica": "PROPUESTA TÉCNICA" # Nueva Sección
            }

            letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
            current_letter_idx = 0

            for key in selected_sections_keys:
                if key == "contenido_separadores": continue # No se lista a sí mismo

                name = readable_names.get(key, key.upper().replace("_", " "))
                
                # Obtener número de página
                page_num = ""
                if page_mapping and key in page_mapping:
                    page_num = f"{page_mapping[key]}"

                p = doc.add_paragraph()
                
                # Letra y Nombre
                run_letter = p.add_run(f"{letters[current_letter_idx]}. {name}")
                run_letter.bold = True
                run_letter.font.size = Pt(14)
                
                # Puntos de relleno y número de página
                if page_num:
                    # Calcular cuántos puntos agregar (simple estimación visual)
                    # Un enfoque mejor sería usar tabulaciones, pero esto es un fix rápido
                    max_chars = 60 # Ajustar según ancho página
                    current_len = len(f"{letters[current_letter_idx]}. {name}")
                    dots_needed = max(5, max_chars - current_len - len(page_num))
                    
                    run_dots = p.add_run(" " + "." * dots_needed + " ")
                    run_dots.font.size = Pt(14)
                    
                    run_page = p.add_run(page_num)
                    run_page.bold = True
                    run_page.font.size = Pt(14)

                current_letter_idx += 1

            # Guardar Word temporal
            temp_word = os.path.abspath("temp_separadores.docx")
            doc.save(temp_word)

            # Convertir a PDF usando Word Automation (igual que ExcelToWord)
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                wb = word.Documents.Open(temp_word)
                wb.SaveAs(output_path, FileFormat=17) # 17 = PDF
                wb.Close()
            finally:
                word.Quit()
            
            # Limpiar temporal
            if os.path.exists(temp_word):
                os.remove(temp_word)

            return True

        except Exception as e:
            print(f"Error generando separadores: {e}")
            return False

    def merge_pdfs(self, output_path, ordered_items, generated_quotation_pdf, external_files_map):
        """
        Une los PDFs en el orden especificado.
        
        :param output_path: Ruta final del PDF unido.
        :param ordered_items: Lista de claves (e.g. ['portadas', 'separadores', 'external_1'])
        :param generated_quotation_pdf: Ruta al PDF de la cotización generado (Presupuesto).
        :param external_files_map: Dict mapping keys (e.g. 'external_1') -> file_path
        """
        merger = PyPDF2.PdfMerger()
        
        try:
            # --- FASE 1: CALCULAR PÁGINAS ---
            current_page = 1
            page_mapping = {}
            
            # Necesitamos saber las rutas antes de unir, para contar páginas
            # Para 'separadores', asumimos 1 página (casi siempre es suficiente)
            
            for item_key in ordered_items:
                pdf_path = None
                pages_in_item = 0
                
                if item_key == "contenido_separadores":
                    # Asumimos que la tabla de contenido será 1 sola página
                    # Esto es necesario porque aún no la hemos generado
                    pdf_path = "PLACEHOLDER_SEPARATORS" 
                    pages_in_item = 1
                    
                elif item_key == "presupuesto_programacion":
                    pdf_path = generated_quotation_pdf
                
                elif item_key in self.SECTION_MAP:
                    filename = self.SECTION_MAP[item_key]
                    pdf_path = os.path.join(self.templates_dir, filename)
                
                elif item_key in external_files_map:
                    pdf_path = external_files_map[item_key]
                    
                # Calcular páginas
                if pdf_path and pdf_path != "PLACEHOLDER_SEPARATORS" and os.path.exists(pdf_path):
                     try:
                         reader = PyPDF2.PdfReader(pdf_path)
                         pages_in_item = len(reader.pages)
                     except Exception as e:
                         print(f"Error leyendo páginas de {item_key}: {e}")
                         pages_in_item = 0
                
                # Guardar en mapa (donde empieza esta sección)
                page_mapping[item_key] = current_page
                current_page += pages_in_item

            # --- FASE 2: GENERAR SEPARADORES CON NÚMEROS ---
            separators_pdf_path = None
            if "contenido_separadores" in ordered_items:
                 separators_pdf_path = os.path.abspath("temp_separadores_final.pdf")
                 # Filtrar solo secciones "reales" para la lista
                 sections_for_toc = [i for i in ordered_items if i in self.SECTION_MAP or i == "presupuesto_programacion" or i == "propuesta_tecnica"]
                 
                 # Pasar el mapa de páginas
                 self.generate_separators_pdf(sections_for_toc, separators_pdf_path, page_mapping)

            # --- FASE 3: UNIR ---
            for item_key in ordered_items:
                pdf_to_add = None
                
                if item_key == "presupuesto_programacion":
                    pdf_to_add = generated_quotation_pdf
                
                elif item_key == "contenido_separadores":
                    pdf_to_add = separators_pdf_path

                elif item_key in self.SECTION_MAP:
                    filename = self.SECTION_MAP[item_key]
                    pdf_to_add = os.path.join(self.templates_dir, filename)
                
                elif item_key in external_files_map:
                    pdf_to_add = external_files_map[item_key]

                # Agregar si existe
                if pdf_to_add and os.path.exists(pdf_to_add):
                    merger.append(pdf_to_add)
                else:
                    print(f"Advertencia: No se encontró PDF para '{item_key}' en {pdf_to_add}")

            merger.write(output_path)
            merger.close()

            # Limpiar separadores temporal
            if separators_pdf_path and os.path.exists(separators_pdf_path):
                os.remove(separators_pdf_path)
                
            return True

        except Exception as e:
            print(f"Error uniendo PDFs: {e}")
            return False
