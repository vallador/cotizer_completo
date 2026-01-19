from docx import Document
import os
import re
from datetime import datetime



class WordController:
    def __init__(self, cotizacion_controller):
        self.cotizacion_controller = cotizacion_controller

        base_path = os.path.dirname(os.path.abspath(__file__))
        self.templates_dir = os.path.join(base_path, 'data', 'templates')

        # Nombres de plantillas que deben existir en el directorio
        self.template_juridica_larga = os.path.join(self.templates_dir, 'plantilla_juridica_larga.docx')
        self.template_juridica_corta = os.path.join(self.templates_dir, 'plantilla_juridica_corta.docx')
        self.template_natural_corta = os.path.join(self.templates_dir, 'plantilla_natural_corta.docx')

        # Verificar que existen las plantillas
        self._verify_templates_exist()

    def _verify_templates_exist(self):
        """Verifica que las plantillas existan"""
        templates = {
            'Jurídica Larga': self.template_juridica_larga,
            'Jurídica Corta': self.template_juridica_corta,
            'Natural Corta': self.template_natural_corta
        }

        missing_templates = []
        for name, path in templates.items():
            if not os.path.exists(path):
                missing_templates.append(f"  - {name}: {path}")

        if missing_templates:
            print("⚠️  ADVERTENCIA: Las siguientes plantillas no se encontraron:")
            for template in missing_templates:
                print(template)
            print(f"\nAsegúrate de que estén en: {self.templates_dir}")

    def generate_word_document(self, cotizacion_id, excel_path, datos_adicionales, formato='auto',
                               template_personalizada=None):
        """
        Genera un documento Word usando las plantillas existentes.

        Args:
            cotizacion_id: ID de la cotización
            excel_path: Ruta al archivo Excel de la cotización
            datos_adicionales: Diccionario con datos adicionales para el documento
            formato: 'largo', 'corto' o 'auto' (selecciona automáticamente según tipo de cliente)
            template_personalizada: Ruta a una plantilla personalizada (opcional)

        Returns:
            Ruta al documento Word generado
        """
        # Obtener la cotización
        cotizacion = self.cotizacion_controller.obtener_cotizacion(cotizacion_id)
        if not cotizacion:
            raise ValueError(f"No se encontró la cotización con ID {cotizacion_id}")

        # Determinar la plantilla a utilizar
        if template_personalizada and os.path.exists(template_personalizada):
            template_path = template_personalizada
        else:
            template_path = self._select_template(cotizacion, formato)

        # Verificar que la plantilla existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")


        # Preparar los datos para reemplazo
        replace_data = self._prepare_replace_data(cotizacion, datos_adicionales)

        # Generar el documento
        output_path = self._generate_document(template_path, replace_data)

        return output_path

    def _select_template(self, cotizacion, formato):
        """
        Selecciona la plantilla adecuada según el tipo de cliente y formato.
        """
        tipo_cliente = cotizacion.cliente.tipo.lower()

        if formato == 'auto':
            if tipo_cliente == 'juridica':
                return self.template_juridica_larga
            else:
                return self.template_natural_corta
        elif formato == 'largo':
            if tipo_cliente == 'juridica':
                return self.template_juridica_larga
            else:
                return self.template_natural_corta
        elif formato == 'corto':
            if tipo_cliente == 'juridica':
                return self.template_juridica_corta
            else:
                return self.template_natural_corta
        else:
            return self._select_template(cotizacion, 'auto')

    def debug_template_markers(self, template_path=None):
        """
        Ver qué marcadores están en una plantilla.
        """
        results = {}

        if template_path:
            if os.path.exists(template_path):
                results[os.path.basename(template_path)] = self._extract_markers_from_template(template_path)
            else:
                results['ERROR'] = f"No se encontró la plantilla: {template_path}"
        else:
            # Analizar todas las plantillas
            templates = {
                'plantilla_juridica_larga.docx': self.template_juridica_larga,
                'plantilla_juridica_corta.docx': self.template_juridica_corta,
                'plantilla_natural_corta.docx': self.template_natural_corta
            }

            for name, path in templates.items():
                if os.path.exists(path):
                    results[name] = self._extract_markers_from_template(path)
                else:
                    results[name] = ['PLANTILLA NO ENCONTRADA']

        return results

    def _extract_markers_from_template(self, template_path):
        """Extrae todos los marcadores {{}} de una plantilla"""
        try:
            doc = Document(template_path)
            markers_found = set()

            # Buscar en párrafos
            for paragraph in doc.paragraphs:
                matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
                markers_found.update(matches)

            # Buscar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
                            markers_found.update(matches)

            # Buscar en encabezados y pies de página
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
                        markers_found.update(matches)

                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
                        markers_found.update(matches)

            return sorted(list(markers_found))

        except Exception as e:
            return [f"ERROR: {str(e)}"]

    def list_available_markers(self):
        """Lista todos los marcadores disponibles"""
        sample_data = {
            'fecha': 'Fecha formateada (ej: 16 de septiembre de 2025)',
            'cliente': 'Nombre del cliente',
            'nit': 'NIT del cliente',
            'direccion': 'Dirección del cliente',
            'referencia': 'Referencia del proyecto',
            'validez': 'Días de validez de la oferta',
            'cuadrillas': 'Número de cuadrillas',
            'operarios': 'Número de operarios',
            'plazo': 'Días de plazo de ejecución',
            'forma_pago': 'Forma de pago (se genera automáticamente según configuración)',
            'tabla_cotizacion': 'Se reemplaza por imagen de la tabla de Excel'
        }

        print("📋 MARCADORES DISPONIBLES PARA USAR EN LAS PLANTILLAS:")
        print("=" * 60)
        for marker, description in sample_data.items():
            print(f"{{{{ {marker:<18} }}}} → {description}")
        print("=" * 60)

        return sample_data


    def _num_to_text(self, number):
        """Convierte números del 1 al 99 a texto en español."""
        try:
            number = int(number)
            if number == 1: return "un" # "Una cuadrilla" suena mejor que "uno cuadrilla"? Ajustar según contexto.
            if number == 1: return "una" if "cuadrilla" in str(number) else "un" # Simplificado
            
            unidades = {1: "un", 2: "dos", 3: "tres", 4: "cuatro", 5: "cinco", 
                        6: "seis", 7: "siete", 8: "ocho", 9: "nueve"}
            decenas = {10: "diez", 11: "once", 12: "doce", 13: "trece", 14: "catorce", 15: "quince",
                       16: "dieciséis", 17: "diecisiete", 18: "dieciocho", 19: "diecinueve",
                       20: "veinte", 30: "treinta", 40: "cuarenta", 50: "cincuenta",
                       60: "sesenta", 70: "setenta", 80: "ochenta", 90: "noventa"}
            
            if number in unidades: return unidades[number]
            if number in decenas: return decenas[number]
            if number < 30: return "veinti" + unidades[number % 10]
            
            decena = (number // 10) * 10
            unidad = number % 10
            if unidad == 0: return decenas[decena]
            return f"{decenas[decena]} y {unidades[unidad]}"
        except:
            return str(number)

    def _prepare_replace_data(self, cotizacion, datos_adicionales):
        """Prepara los datos para reemplazar en la plantilla."""
        fecha_actual = datetime.now()
        meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                 'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
        fecha_formateada = f"{fecha_actual.day} de {meses[fecha_actual.month - 1]} de {fecha_actual.year}"

        # --- 1. Lógica de Pago ---
        forma_pago = "Contraentrega" # Default
        texto_forma_pago = None # Initialize variable

        if datos_adicionales.get('pago_contraentrega'):
            forma_pago = "Contraentrega"
        elif datos_adicionales.get('pago_porcentajes'):
            anticipo = int(datos_adicionales.get('anticipo', 0))
            anticipo_text = self._num_to_text(anticipo)
            
            avance_req = datos_adicionales.get('avance_requerido', 50)
            avance_pago = datos_adicionales.get('avance', 0)
            final = datos_adicionales.get('final', 0)
            
            texto_forma_pago = (
                f"Un anticipo en la cuantía equivalente al {anticipo_text} por ciento ({anticipo}%) del valor del Contrato "
                f"el {anticipo}% en cortes de obra y el {100-anticipo}% restante como a la terminación del contrato garantizando la calidad del trabajo, "
                f"el anticipo será exclusivamente para el alistamiento de la iniciación de la obra y para la compra y sostenimiento de precios de productos y materiales. "
                f"Se harán pagos parciales que estarán soportados por un Acta de avance de obra firmada entre el Administrador o persona encargada de la obra y el residente de obra designado por el contratista."
            )
        elif datos_adicionales.get('pago_personalizado'):
             forma_pago = datos_adicionales.get('pago_personalizado', '')

        # --- 2. Lógica de Pólizas (Estructurada) ---
        polizas_textos = {
            "manejo_anticipo": ("A.) BUEN MANEJO DEL ANTICIPO Y CORRECTA INVERSIÓN: ", "Por un valor equivalente al Ciento por ciento (100%) del valor entregado como anticipo de la obra al Contratista, con una vigencia igual al plazo de ejecución del Contrato y seis (6) meses más, con fecha de expedición a partir de la entrega del anticipo."),
            "cumplimiento_contrato": ("B.) CUMPLIMIENTO DEL CONTRATO: ", "Por un valor equivalente al treinta por ciento (30%) del valor total del Contrato, para garantizar el cumplimiento de las obligaciones contraídas en virtud del Contrato, con una vigencia igual al plazo de ejecución del Contrato y seis (6) meses más, con fecha de expedición a partir de la suscripción del Contrato."),
            "calidad_servicio": ("C.) CALIDAD DE SERVICIO: ", "Por un valor equivalente al Diez por ciento (10%) del valor final del Contrato, con una vigencia igual a dos (2) años contados a partir de la fecha de la suscripción del acta de recibo final de la obra a satisfacción."),
            "pago_salarios": ("D.) PAGO DE SALARIOS, PRESTACIONES SOCIALES E INDEMNIZACIONES LABORALES DEL PERSONAL QUE PRESTE SUS SERVICIOS EN LA EJECUCIÓN DEL CONTRATO: ", "Por un valor equivalente al Quince por ciento (15%) del valor total del Contrato, con una vigencia igual al plazo de ejecución del Contrato y tres (3) años más."),
            "responsabilidad_civil": ("E.) SEGURO DE RESPONSABILIDAD CIVIL EXTRACONTRACTUAL: ", "Se debe constituir dentro de los tres (3) días siguientes a la firma del Contrato, por un valor equivalente al Diez por ciento (10%) del valor total del Contrato, con una vigencia igual al plazo de ejecución del Contrato y tres (3) meses más."),
            "calidad_funcionamiento": ("F.) CALIDAD Y CORRECTO FUNCIONAMIENTO: ", "Para garantizar la calidad y eficiencia de los equipos suministrados, así como las obras civiles requeridas, por valor del 20% del total resultante del contrato y con una vigencia durante la garantía de los equipos y un (1) año adicional.")
        }
        
        polizas_seleccionadas = datos_adicionales.get('polizas_incluir', {})
        lista_polizas = []
        for key, (titulo, cuerpo) in polizas_textos.items():
            if polizas_seleccionadas.get(key, False):
                lista_polizas.append({'title': titulo, 'body': cuerpo})

        if not lista_polizas:
             # Si no hay, mandamos un string simple
             polizas_data = "No requiere pólizas específicas."
        else:
             polizas_data = lista_polizas

        # --- 3. Preparación de Datos ---
        cuadrillas = int(datos_adicionales.get('cuadrillas', 1))
        operarios = int(datos_adicionales.get('operarios', 2))
        plazo_dias = datos_adicionales.get('plazo_dias', 15)
        plazo_tipo = datos_adicionales.get('plazo_tipo', 'calendario')

        # Si el valor de 'operarios_letra' viene vacío, lo calculamos
        operarios_texto = datos_adicionales.get('operarios_letra', '')
        if not operarios_texto or operarios_texto == 'dos': 
             operarios_texto = self._num_to_text(operarios)

        replace_data = {
            'fecha': fecha_formateada,
            'cliente': cotizacion.cliente.nombre,
            'nit': cotizacion.cliente.nit or 'N/A',
            'direccion': cotizacion.cliente.direccion or 'N/A',
            'referencia': datos_adicionales.get('referencia', ''),
            'validez_oferta': str(datos_adicionales.get('validez', '30')),
            
            # Personal
            'cuadrillas_numero': str(cuadrillas),
            'cuadrillas_texto': self._num_to_text(cuadrillas),
            'operarios_numero': str(operarios),
            'operarios_texto': operarios_texto,
            
            # Plazo y Pago
            'plazo_texto': f"{plazo_dias} días {plazo_tipo}",
            'forma_pago': forma_pago if forma_pago != "Contraentrega" else "Pago del 100% recibiendo todo a conformidad.",
            'polizas': polizas_data,
            
            # Disponibilidad Personal
            'director_disp': datos_adicionales.get('director_obra', 'N/A'),
            'residente_disp': datos_adicionales.get('residente_obra', 'N/A'),
            'tecnologo_disp': datos_adicionales.get('tecnologo', 'N/A'), 
        }

        # Override payment text if it was generated above
        if 'texto_forma_pago' in locals():
            replace_data['forma_pago'] = texto_forma_pago

        return replace_data

    def _generate_document(self, template_path, replace_data):
        """
        Genera el documento Word reemplazando los marcadores.
        """
        try:
            # Cargar la plantilla
            doc = Document(template_path)

            # Reemplazo principal
            for key, value in replace_data.items():
                marker = f"{{{{{key}}}}}"
                
                # Buscar en párrafos
                for paragraph in doc.paragraphs:
                    if marker in paragraph.text:
                         self._replace_text_in_paragraph(paragraph, marker, value)

                # Buscar en tablas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if marker in paragraph.text:
                                    self._replace_text_in_paragraph(paragraph, marker, value)
                                    
                # Buscar en secciones (header/footer)
                for section in doc.sections:
                    if section.header:
                         for p in section.header.paragraphs:
                             if marker in p.text: self._replace_text_in_paragraph(p, marker, value)
                    if section.footer:
                         for p in section.footer.paragraphs:
                             if marker in p.text: self._replace_text_in_paragraph(p, marker, value)

            # Generar archivo de salida
            output_dir = os.path.join(os.getcwd(), 'output')
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            output_filename = f"Cotizacion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(output_dir, output_filename)

            doc.save(output_path)
            return output_path

        except Exception as e:
            print(f"Error al generar el documento Word: {e}")
            raise

    def _replace_text_in_paragraph(self, paragraph, marker, content):
        """
        Reemplaza el marcador e impone negrita al contenido insertado.
        Si content es list: Reemplazo estructurado (Polizas).
        Si content es str: Reemplazo en linea con reconstruccion de runs.
        """
        if isinstance(content, list):
             # Lógica para Pólizas (Lista de dicts)
             paragraph.clear()
             
             for item in content:
                 title = item.get('title', '')
                 body = item.get('body', '')
                 
                 # Título (Negrita)
                 run_title = paragraph.add_run(title)
                 run_title.bold = True
                 
                 # Cuerpo (AHORA TAMBIÉN NEGRISTA según solicitud "todo lo que agregue... en negrita")
                 run_body = paragraph.add_run(body)
                 run_body.bold = True
                 
                 # Salto de línea
                 paragraph.add_run("\n\n")
                 
        else:
            # Lógica estándar (String) con Negrita Forzada
            if marker not in paragraph.text:
                return

            text_content = str(content)
            
            # --- Revisar si el marcador está "partido" entre runs (Split Token problem) ---
            # Si el marcador está en el párrafo pero no en ningún run individual,
            # limpiamos el párrafo y hacemos un reemplazo simple.
            marker_is_split = True
            for run in paragraph.runs:
                if marker in run.text:
                    marker_is_split = False
                    break
            
            if marker_is_split:
                # Estrategia de Fallback: Reemplazo simple (perdemos formato mixto, pero garantizamos el valor)
                full_text = paragraph.text.replace(marker, text_content)
                
                # Intentar conservar el estilo del primer run si existe
                first_run_style = {}
                if paragraph.runs:
                    r1 = paragraph.runs[0]
                    first_run_style = {
                         'name': r1.font.name,
                         'size': r1.font.size,
                         'bold': r1.bold,
                         'italic': r1.italic
                    }
                
                paragraph.clear()
                run = paragraph.add_run(full_text)
                
                # Restaurar estilo base
                if first_run_style.get('name'): run.font.name = first_run_style['name']
                if first_run_style.get('size'): run.font.size = first_run_style['size']
                # run.bold = first_run_style.get('bold') # No usamos el original, queremos controlar el bold?
                
                # INTENTO DE RESALTAR: Como no sabemos dónde quedó el texto insertado exactamente en el string plano,
                # y el usuario quiere "lo que se agrega en negrita", en este fallback es difícil.
                # Opción: Dejarlo como está (formato del primer run) y advertir.
                # O mejor: Asumir que si está partido, probablemente es un header completo.
                
                # IMPORTANTE: Si hacemos este fallback, el valor APARECE, que es lo más importante ahora.
                return

            # --- Caso Normal: El marcador existe completo en al menos un run ---
            # 1. Capturar estado actual de los runs
            runs_data = []
            for run in paragraph.runs:
                runs_data.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'color': run.font.color.rgb if run.font.color else None,
                    'name': run.font.name,
                    'size': run.font.size
                })

            # 2. Limpiar párrafo
            paragraph.clear()

            # 3. Reconstruir
            for data in runs_data:
                text = data['text']
                if marker in text:
                    # Dividir el run
                    parts = text.split(marker)
                    
                    for i, part in enumerate(parts):
                        # Agregar parte texto original
                        if part:
                            r = paragraph.add_run(part)
                            self._copy_run_style(r, data)
                        
                        # Agregar marcador (si no es el último fragmento)
                        if i < len(parts) - 1:
                            r_new = paragraph.add_run(text_content)
                            self._copy_run_style(r_new, data)
                            r_new.bold = True # FORZAR NEGRITA
                else:
                    # Run normal sin marcador
                    r = paragraph.add_run(text)
                    self._copy_run_style(r, data)

    def _copy_run_style(self, new_run, style_data):
        """Copia estilos básicos al nuevo run."""
        new_run.bold = style_data['bold']
        new_run.italic = style_data['italic']
        new_run.underline = style_data['underline']
        if style_data['name']:
            new_run.font.name = style_data['name']
        if style_data['size']:
            new_run.font.size = style_data['size']
        if style_data['color']:
             new_run.font.color.rgb = style_data['color']